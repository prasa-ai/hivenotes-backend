"""
LangGraph node — Step 7: Generate a Microsoft DOCX file from the SOAP sections
produced in Step 6.

Document layout
───────────────
  ┌─────────────────────────────────────────┐
  │  SOAP Therapy Note                      │  ← Title
  │  Client: <client_id>                    │
  │  Therapist: <therapist_id>              │  ← Metadata table
  │  Session: <session_id>                  │
  │  Generated: <UTC timestamp>             │
  ├─────────────────────────────────────────┤
  │  Subjective                             │  ← Heading 1
  │  <text>                                 │
  │  Objective                              │  ← Heading 1
  │  <text>                                 │
  │  Assessment                             │  ← Heading 1
  │  <text>                                 │
  │  Plan                                   │  ← Heading 1
  │  <text>                                 │
  └─────────────────────────────────────────┘

Reads:   state["soap_sections"], state["therapist_id"], state["client_id"],
         state["session_id"]
Writes:  state["docx_bytes"]  — in-memory .docx bytes, passed to Step 8
"""
import io
import logging
from datetime import datetime, timezone
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.workflow.state import GraphState

logger = logging.getLogger(__name__)

# ── Style constants ────────────────────────────────────────────────────────────
_BRAND_BLUE   = RGBColor(0x1A, 0x73, 0xE8)   # title colour
_SECTION_COLOUR = RGBColor(0x20, 0x20, 0x20)  # heading colour
_SOAP_SECTIONS = [
    ("Subjective",  "subjective"),
    ("Objective",   "objective"),
    ("Assessment",  "assessment"),
    ("Plan",        "plan"),
]


def _build_document(
    soap_sections: dict,
    therapist_id: str,
    client_id: str,
    session_id: str,
) -> Document:
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Title ─────────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("SOAP Therapy Note")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = _BRAND_BLUE

    doc.add_paragraph()  # spacer

    # ── Metadata table ────────────────────────────────────────────────────────
    generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    meta_rows = [
        ("Therapist",  therapist_id),
        ("Client",     client_id),
        ("Session",    session_id),
        ("Generated",  generated_at),
    ]
    meta_table = doc.add_table(rows=len(meta_rows), cols=2)
    meta_table.style = "Table Grid"
    for i, (label, value) in enumerate(meta_rows):
        label_cell = meta_table.cell(i, 0)
        value_cell = meta_table.cell(i, 1)
        # Label cell — bold
        label_run = label_cell.paragraphs[0].add_run(label)
        label_run.bold = True
        label_run.font.size = Pt(10)
        # Value cell
        value_run = value_cell.paragraphs[0].add_run(value)
        value_run.font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # ── SOAP sections ─────────────────────────────────────────────────────────
    for heading_label, state_key in _SOAP_SECTIONS:
        section_text: str = soap_sections.get(state_key, "Not documented in this session.")

        # Section heading
        heading = doc.add_heading(heading_label, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in heading.runs:
            run.font.color.rgb = _SECTION_COLOUR
            run.font.size = Pt(13)

        # Section body — respect line breaks from the model output
        for line in section_text.splitlines():
            line = line.strip()
            if line:
                body_para = doc.add_paragraph(line)
                body_para.style = "Normal"
                for run in body_para.runs:
                    run.font.size = Pt(11)
            else:
                doc.add_paragraph()  # blank line between natural paragraphs

        doc.add_paragraph()  # spacer between sections

    # ── Footer ────────────────────────────────────────────────────────────────
    footer_para = doc.sections[0].footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        f"Confidential — SOAP Notes AI  |  Session {session_id}  |  {generated_at}"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    return doc


async def generate_docx_node(state: GraphState) -> GraphState:
    """
    Build a formatted .docx from the SOAP sections dict and store the raw
    bytes in state["docx_bytes"] for upload in Step 8.
    """
    soap_sections: dict = state.get("soap_sections", {})
    therapist_id: str   = state.get("therapist_id", "unknown")
    client_id: str      = state.get("client_id", "unknown")
    session_id: str     = state.get("session_id", "unknown")

    if not soap_sections:
        msg = "generate_docx: soap_sections is empty — generate_soap node may have failed."
        logger.error(msg)
        return {**state, "error": msg}

    logger.info(
        "generate_docx: building DOCX for therapist=%s client=%s session=%s",
        therapist_id, client_id, session_id,
    )

    try:
        doc = _build_document(soap_sections, therapist_id, client_id, session_id)

        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()

    except Exception as exc:
        msg = f"generate_docx: failed to build document — {exc}"
        logger.error(msg)
        return {**state, "error": msg}

    logger.info("generate_docx: DOCX built — %d bytes", len(docx_bytes))

    return {
        **state,
        "docx_bytes": docx_bytes,
        "error":      None,
    }
